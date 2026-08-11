#!/usr/bin/env python3
import json
import sys
from pathlib import Path

from docx import Document


def main() -> None:
    source = Path(sys.argv[1])
    output = Path(sys.argv[2])
    doc = Document(source)
    blocks = []
    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text:
            blocks.append({
                "type": "paragraph",
                "index": index,
                "style": paragraph.style.name if paragraph.style else None,
                "text": text,
            })
    for table_index, table in enumerate(doc.tables):
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        blocks.append({"type": "table", "index": table_index, "rows": rows})
    payload = {
        "source_document": source.name,
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "blocks": blocks,
    }
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


if __name__ == "__main__":
    main()
